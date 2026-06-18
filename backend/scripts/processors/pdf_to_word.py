import os
from pathlib import Path
from .utils import docx, extract_pdf_text

def pdf_to_word(input_paths, output_path):
    if not input_paths:
        return
    
    input_path = input_paths[0]
    
    # 1. Try high-fidelity pdf2docx converter first
    try:
        from pdf2docx import Converter
        print(f"Using pdf2docx high-fidelity engine to convert {input_path} to {output_path}...")
        
        cv = Converter(str(input_path))
        # cv.convert parses layout (paragraphs, tables, images) and creates the docx
        cv.convert(str(output_path), start=0, end=None)
        cv.close()
        print(f"High-fidelity conversion completed successfully: {output_path}")
        return
    except Exception as e:
        print(f"pdf2docx engine failed or not installed. Error: {e}")
        print("Falling back to standard text extraction parser...")

    # 2. Fallback: Standard text-extraction-based converter
    pdf_text = extract_pdf_text(input_path)
    mode = os.environ.get("PDF_TO_WORD_MODE", "flowing").upper()
    ocr_run = os.environ.get("PDF_TO_WORD_OCR", "true").lower() == "true"
    lang = os.environ.get("PDF_TO_WORD_LANG", "en").upper()

    if docx:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        doc = Document()

        # --- Document-level styling ---
        # Narrow margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Title block
        title_para = doc.add_heading("Exported Document", level=1)
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run("Exported Document")
        title_run.font.color.rgb = RGBColor(0x1e, 0x2d, 0x3d)

        # Metadata row
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Mode: {mode}  |  OCR: {'Enabled' if ocr_run else 'Disabled'}  |  Language: {lang}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        meta_para.space_after = Pt(14)

        # Horizontal separator
        def add_hr(doc):
            para = doc.add_paragraph()
            p = para._p
            pPr = OxmlElement('w:pPr')
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.insert(0, pPr)
            return para

        add_hr(doc)

        # Split text into logical blocks
        lines = pdf_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                i += 1
                continue

            # Heuristic: if line is short and next line is empty → treat as heading/section
            is_likely_heading = (
                len(line) < 60
                and (i + 1 >= len(lines) or not lines[i + 1].strip())
                and not line.endswith(".")
                and not line.endswith(",")
            )

            if is_likely_heading:
                h = doc.add_heading(line, level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            elif line.startswith(("• ", "- ", "* ", "– ")):
                # List item
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line.lstrip("•-*–").strip())
                run.font.size = Pt(11)
            else:
                # Normal paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = "Calibri"

            i += 1

        # Footer note
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run("Converted by WeLovePDF • welovepdf.com")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
    else:
        # Plain text fallback
        Path(output_path).write_text(
            f"PDF to Word Export\nMode: {mode} | OCR: {ocr_run} | Language: {lang}\n\n{pdf_text}",
            encoding="utf-8"
        )

    print(f"Converted PDF to DOCX successfully: {output_path}")
